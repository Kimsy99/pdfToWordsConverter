import pdfplumber
from docx import Document
import os
from multiprocessing import Process  # to improve converting speed


def convertPdf(fileName):
    with pdfplumber.open(fileName) as pdf:
        print("Converting pdf file to {0}.words total {1} pages.".format(
            fileName, len(pdf.pages)))
        content = ''
        baseName = fileName.split('.')[0]
        wordName = baseName + '.docx'
        flag = True
        if os.path.exists(wordName):
            os.remove(wordName)
        for i in range(len(pdf.pages)):
            print("Converting {0}.pdf page {1}".format(baseName, i))
            page = pdf.pages[i]
            if page.extract_text() == None:
                print("No text to retrive on page {0}".format(fileName))
                flag = False
                break
            page_content = '\n'.join(page.extract_text().split('\n')[:-1])
            content = content + page_content
            if os.path.exists(wordName):
                doc = Document(wordName)
            else:
                doc = Document()
            doc.add_paragraph(content)
            doc.save(wordName)
            content = ''
            print('Page {0} completed'.format(i))
        if flag:
            print('Completed converting to {}.docx'.format(baseName))


if __name__ == '__main__':
    for file in os.listdir('.'):
        if os.path.isfile(file) and file.split('.')[1] == 'pdf':
            p = Process(target=convertPdf, args=(file,))
            p.start()
